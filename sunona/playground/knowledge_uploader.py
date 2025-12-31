"""
Sunona Voice AI - Knowledge Uploader

Drag-and-drop document upload for knowledge bases.
Supports multiple file formats and sources.

Features:
- Multi-format upload (PDF, DOCX, TXT, CSV, JSON)
- Website crawling
- Progress tracking
- Chunking configuration
- Vector store integration
"""

import asyncio
import hashlib
import json
import logging
import mimetypes
import os
import uuid
from dataclasses import dataclass, field
from datetime import datetime, timezone
from enum import Enum
from pathlib import Path
from typing import Any, Callable, Dict, List, Optional, Awaitable

logger = logging.getLogger(__name__)


class UploadStatus(Enum):
    """Document upload status."""
    PENDING = "pending"
    UPLOADING = "uploading"
    PROCESSING = "processing"
    CHUNKING = "chunking"
    EMBEDDING = "embedding"
    COMPLETE = "complete"
    FAILED = "failed"


class DocumentType(Enum):
    """Supported document types."""
    PDF = "pdf"
    DOCX = "docx"
    TXT = "txt"
    CSV = "csv"
    JSON = "json"
    HTML = "html"
    MARKDOWN = "markdown"
    URL = "url"


MIME_TYPE_MAP = {
    "application/pdf": DocumentType.PDF,
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": DocumentType.DOCX,
    "text/plain": DocumentType.TXT,
    "text/csv": DocumentType.CSV,
    "application/json": DocumentType.JSON,
    "text/html": DocumentType.HTML,
    "text/markdown": DocumentType.MARKDOWN,
}


@dataclass
class UploadedDocument:
    """A document being uploaded."""
    id: str
    filename: str
    document_type: DocumentType
    size_bytes: int
    status: UploadStatus = UploadStatus.PENDING
    
    # Processing
    content_hash: Optional[str] = None
    chunk_count: int = 0
    error_message: Optional[str] = None
    
    # Progress
    progress_percent: float = 0.0
    
    # Metadata
    knowledge_base_id: Optional[str] = None
    organization_id: Optional[str] = None
    uploaded_by: Optional[str] = None
    created_at: datetime = field(default_factory=lambda: datetime.now(timezone.utc))
    completed_at: Optional[datetime] = None
    
    def to_dict(self) -> Dict[str, Any]:
        return {
            "id": self.id,
            "filename": self.filename,
            "document_type": self.document_type.value,
            "size_bytes": self.size_bytes,
            "status": self.status.value,
            "chunk_count": self.chunk_count,
            "progress_percent": round(self.progress_percent, 1),
            "error_message": self.error_message,
            "created_at": self.created_at.isoformat(),
            "completed_at": self.completed_at.isoformat() if self.completed_at else None,
        }


@dataclass
class ChunkConfig:
    """Configuration for document chunking."""
    chunk_size: int = 1000  # Characters per chunk
    chunk_overlap: int = 200  # Overlap between chunks
    separator: str = "\n\n"  # Preferred split point
    min_chunk_size: int = 100
    
    def to_dict(self) -> Dict[str, int]:
        return {
            "chunk_size": self.chunk_size,
            "chunk_overlap": self.chunk_overlap,
            "min_chunk_size": self.min_chunk_size,
        }


@dataclass
class DocumentChunk:
    """A chunk of a document."""
    id: str
    document_id: str
    content: str
    chunk_index: int
    
    # Metadata
    page_number: Optional[int] = None
    section: Optional[str] = None
    
    # Embedding
    embedding: Optional[List[float]] = None
    
    def to_dict(self) -> Dict[str, Any]:
        return {
            "id": self.id,
            "document_id": self.document_id,
            "content": self.content[:200] + "..." if len(self.content) > 200 else self.content,
            "chunk_index": self.chunk_index,
            "page_number": self.page_number,
            "section": self.section,
            "has_embedding": self.embedding is not None,
        }


class KnowledgeUploader:
    """
    Knowledge base document uploader.
    
    Handles document upload, processing, chunking, and
    embedding for knowledge base creation.
    
    Example:
        uploader = KnowledgeUploader()
        
        # Upload a file
        doc = await uploader.upload_file(
            file_data=pdf_bytes,
            filename="manual.pdf",
            knowledge_base_id="kb_123",
        )
        
        # Check progress
        status = await uploader.get_status(doc.id)
        
        # Upload from URL
        doc = await uploader.upload_url(
            url="https://example.com/docs",
            knowledge_base_id="kb_123",
        )
    """
    
    def __init__(
        self,
        storage_path: str = "uploads",
        vector_store: Optional[Any] = None,
        embedding_service: Optional[Any] = None,
    ):
        """
        Initialize knowledge uploader.
        
        Args:
            storage_path: Directory for storing uploads
            vector_store: Vector store for embeddings
            embedding_service: Service for generating embeddings
        """
        self.storage_path = Path(storage_path)
        self.storage_path.mkdir(parents=True, exist_ok=True)
        
        self.vector_store = vector_store
        self.embedding_service = embedding_service
        
        self._documents: Dict[str, UploadedDocument] = {}
        self._chunks: Dict[str, List[DocumentChunk]] = {}  # doc_id -> chunks
        self._processing_tasks: Dict[str, asyncio.Task] = {}
    
    async def upload_file(
        self,
        file_data: bytes,
        filename: str,
        knowledge_base_id: str,
        organization_id: Optional[str] = None,
        uploaded_by: Optional[str] = None,
        chunk_config: Optional[ChunkConfig] = None,
    ) -> UploadedDocument:
        """
        Upload a file for processing.
        
        Args:
            file_data: File content as bytes
            filename: Original filename
            knowledge_base_id: Target knowledge base
            organization_id: Organization ID
            uploaded_by: User ID
            chunk_config: Chunking configuration
        
        Returns:
            UploadedDocument with upload status
        """
        doc_id = f"doc_{uuid.uuid4().hex[:12]}"
        
        # Detect document type
        mime_type, _ = mimetypes.guess_type(filename)
        doc_type = MIME_TYPE_MAP.get(mime_type)
        
        if not doc_type:
            ext = Path(filename).suffix.lower().lstrip(".")
            try:
                doc_type = DocumentType(ext)
            except ValueError:
                doc_type = DocumentType.TXT
        
        # Create document record
        doc = UploadedDocument(
            id=doc_id,
            filename=filename,
            document_type=doc_type,
            size_bytes=len(file_data),
            status=UploadStatus.UPLOADING,
            content_hash=hashlib.sha256(file_data).hexdigest()[:16],
            knowledge_base_id=knowledge_base_id,
            organization_id=organization_id,
            uploaded_by=uploaded_by,
        )
        
        self._documents[doc_id] = doc
        
        # Save file
        file_path = self.storage_path / f"{doc_id}_{filename}"
        with open(file_path, "wb") as f:
            f.write(file_data)
        
        doc.status = UploadStatus.PROCESSING
        doc.progress_percent = 10
        
        # Start async processing
        config = chunk_config or ChunkConfig()
        task = asyncio.create_task(
            self._process_document(doc, file_path, config)
        )
        self._processing_tasks[doc_id] = task
        
        logger.info(f"Uploaded document {doc_id}: {filename}")
        return doc
    
    async def upload_url(
        self,
        url: str,
        knowledge_base_id: str,
        organization_id: Optional[str] = None,
        uploaded_by: Optional[str] = None,
        max_pages: int = 10,
        chunk_config: Optional[ChunkConfig] = None,
    ) -> UploadedDocument:
        """
        Upload content from a URL.
        
        Args:
            url: URL to crawl
            knowledge_base_id: Target knowledge base
            organization_id: Organization ID
            uploaded_by: User ID
            max_pages: Maximum pages to crawl
            chunk_config: Chunking configuration
        
        Returns:
            UploadedDocument with upload status
        """
        doc_id = f"doc_{uuid.uuid4().hex[:12]}"
        
        doc = UploadedDocument(
            id=doc_id,
            filename=url,
            document_type=DocumentType.URL,
            size_bytes=0,
            status=UploadStatus.PROCESSING,
            knowledge_base_id=knowledge_base_id,
            organization_id=organization_id,
            uploaded_by=uploaded_by,
        )
        
        self._documents[doc_id] = doc
        
        # Start async crawling
        config = chunk_config or ChunkConfig()
        task = asyncio.create_task(
            self._crawl_url(doc, url, max_pages, config)
        )
        self._processing_tasks[doc_id] = task
        
        logger.info(f"Started URL crawl {doc_id}: {url}")
        return doc
    
    async def _process_document(
        self,
        doc: UploadedDocument,
        file_path: Path,
        config: ChunkConfig,
    ) -> None:
        """Process an uploaded document."""
        try:
            # Extract text
            doc.progress_percent = 20
            text = await self._extract_text(doc.document_type, file_path)
            
            if not text:
                raise ValueError("Could not extract text from document")
            
            # Chunk text
            doc.status = UploadStatus.CHUNKING
            doc.progress_percent = 50
            chunks = self._chunk_text(doc.id, text, config)
            
            self._chunks[doc.id] = chunks
            doc.chunk_count = len(chunks)
            
            # Generate embeddings
            if self.embedding_service and self.vector_store:
                doc.status = UploadStatus.EMBEDDING
                doc.progress_percent = 70
                await self._embed_chunks(chunks)
                doc.progress_percent = 90
            
            # Complete
            doc.status = UploadStatus.COMPLETE
            doc.progress_percent = 100
            doc.completed_at = datetime.now(timezone.utc)
            
            logger.info(
                f"Processed document {doc.id}: "
                f"{doc.chunk_count} chunks from {doc.filename}"
            )
            
        except Exception as e:
            doc.status = UploadStatus.FAILED
            doc.error_message = str(e)
            logger.error(f"Failed to process document {doc.id}: {e}")
    
    async def _crawl_url(
        self,
        doc: UploadedDocument,
        url: str,
        max_pages: int,
        config: ChunkConfig,
    ) -> None:
        """Crawl a URL and extract content."""
        try:
            import httpx
            
            doc.progress_percent = 10
            
            async with httpx.AsyncClient(timeout=30) as client:
                response = await client.get(url)
                response.raise_for_status()
                
                html = response.text
                doc.size_bytes = len(html.encode())
            
            # Extract text from HTML
            doc.progress_percent = 30
            text = self._extract_text_from_html(html)
            
            # Chunk and embed
            doc.status = UploadStatus.CHUNKING
            doc.progress_percent = 50
            chunks = self._chunk_text(doc.id, text, config)
            
            self._chunks[doc.id] = chunks
            doc.chunk_count = len(chunks)
            
            if self.embedding_service and self.vector_store:
                doc.status = UploadStatus.EMBEDDING
                doc.progress_percent = 70
                await self._embed_chunks(chunks)
            
            doc.status = UploadStatus.COMPLETE
            doc.progress_percent = 100
            doc.completed_at = datetime.now(timezone.utc)
            
            logger.info(f"Crawled URL {doc.id}: {doc.chunk_count} chunks")
            
        except Exception as e:
            doc.status = UploadStatus.FAILED
            doc.error_message = str(e)
            logger.error(f"Failed to crawl URL {doc.id}: {e}")
    
    async def _extract_text(
        self,
        doc_type: DocumentType,
        file_path: Path,
    ) -> Optional[str]:
        """Extract text from a document."""
        try:
            if doc_type == DocumentType.TXT:
                return file_path.read_text(encoding="utf-8")
            
            elif doc_type == DocumentType.MARKDOWN:
                return file_path.read_text(encoding="utf-8")
            
            elif doc_type == DocumentType.JSON:
                data = json.loads(file_path.read_text(encoding="utf-8"))
                return json.dumps(data, indent=2)
            
            elif doc_type == DocumentType.CSV:
                return file_path.read_text(encoding="utf-8")
            
            elif doc_type == DocumentType.PDF:
                return await self._extract_pdf(file_path)
            
            elif doc_type == DocumentType.DOCX:
                return await self._extract_docx(file_path)
            
            elif doc_type == DocumentType.HTML:
                html = file_path.read_text(encoding="utf-8")
                return self._extract_text_from_html(html)
            
            else:
                return file_path.read_text(encoding="utf-8")
                
        except Exception as e:
            logger.error(f"Text extraction error: {e}")
            return None
    
    async def _extract_pdf(self, file_path: Path) -> str:
        """Extract text from PDF."""
        try:
            import pypdf
            
            reader = pypdf.PdfReader(str(file_path))
            text_parts = []
            
            for page in reader.pages:
                text_parts.append(page.extract_text() or "")
            
            return "\n\n".join(text_parts)
        except ImportError:
            logger.warning("pypdf not installed for PDF extraction")
            return ""
    
    async def _extract_docx(self, file_path: Path) -> str:
        """Extract text from DOCX."""
        try:
            from docx import Document
            
            doc = Document(str(file_path))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            
            return "\n\n".join(paragraphs)
        except ImportError:
            logger.warning("python-docx not installed for DOCX extraction")
            return ""
    
    def _extract_text_from_html(self, html: str) -> str:
        """Extract text from HTML."""
        try:
            from bs4 import BeautifulSoup
            
            soup = BeautifulSoup(html, "html.parser")
            
            # Remove scripts and styles
            for tag in soup(["script", "style", "nav", "footer", "header"]):
                tag.decompose()
            
            return soup.get_text(separator="\n", strip=True)
        except ImportError:
            # Simple fallback
            import re
            text = re.sub(r"<[^>]+>", " ", html)
            return " ".join(text.split())
    
    def _chunk_text(
        self,
        doc_id: str,
        text: str,
        config: ChunkConfig,
    ) -> List[DocumentChunk]:
        """Split text into chunks."""
        chunks = []
        
        # Split by separator first
        sections = text.split(config.separator)
        
        current_chunk = ""
        chunk_index = 0
        
        for section in sections:
            if len(current_chunk) + len(section) <= config.chunk_size:
                current_chunk += section + config.separator
            else:
                if len(current_chunk) >= config.min_chunk_size:
                    chunks.append(DocumentChunk(
                        id=f"chunk_{uuid.uuid4().hex[:8]}",
                        document_id=doc_id,
                        content=current_chunk.strip(),
                        chunk_index=chunk_index,
                    ))
                    chunk_index += 1
                
                # Start new chunk with overlap
                overlap_start = max(0, len(current_chunk) - config.chunk_overlap)
                current_chunk = current_chunk[overlap_start:] + section + config.separator
        
        # Add final chunk
        if len(current_chunk.strip()) >= config.min_chunk_size:
            chunks.append(DocumentChunk(
                id=f"chunk_{uuid.uuid4().hex[:8]}",
                document_id=doc_id,
                content=current_chunk.strip(),
                chunk_index=chunk_index,
            ))
        
        return chunks
    
    async def _embed_chunks(self, chunks: List[DocumentChunk]) -> None:
        """Generate embeddings for chunks."""
        if not self.embedding_service:
            return
        
        for chunk in chunks:
            try:
                embedding = await self.embedding_service.embed(chunk.content)
                chunk.embedding = embedding
                
                if self.vector_store:
                    await self.vector_store.add(
                        id=chunk.id,
                        text=chunk.content,
                        embedding=embedding,
                        metadata={
                            "document_id": chunk.document_id,
                            "chunk_index": chunk.chunk_index,
                        },
                    )
            except Exception as e:
                logger.error(f"Embedding error for chunk {chunk.id}: {e}")
    
    async def get_status(self, doc_id: str) -> Optional[UploadedDocument]:
        """Get document upload status."""
        return self._documents.get(doc_id)
    
    async def get_chunks(self, doc_id: str) -> List[DocumentChunk]:
        """Get chunks for a document."""
        return self._chunks.get(doc_id, [])
    
    async def delete_document(self, doc_id: str) -> bool:
        """Delete a document and its chunks."""
        if doc_id not in self._documents:
            return False
        
        doc = self._documents.pop(doc_id)
        self._chunks.pop(doc_id, None)
        
        # Cancel processing if still running
        task = self._processing_tasks.pop(doc_id, None)
        if task and not task.done():
            task.cancel()
        
        # Delete file
        for f in self.storage_path.glob(f"{doc_id}_*"):
            f.unlink()
        
        logger.info(f"Deleted document {doc_id}")
        return True
    
    async def list_documents(
        self,
        knowledge_base_id: Optional[str] = None,
        organization_id: Optional[str] = None,
    ) -> List[UploadedDocument]:
        """List uploaded documents."""
        docs = list(self._documents.values())
        
        if knowledge_base_id:
            docs = [d for d in docs if d.knowledge_base_id == knowledge_base_id]
        
        if organization_id:
            docs = [d for d in docs if d.organization_id == organization_id]
        
        return docs
    
    def get_stats(self) -> Dict[str, Any]:
        """Get uploader statistics."""
        total_chunks = sum(len(c) for c in self._chunks.values())
        
        by_status = {}
        for doc in self._documents.values():
            status = doc.status.value
            by_status[status] = by_status.get(status, 0) + 1
        
        return {
            "total_documents": len(self._documents),
            "total_chunks": total_chunks,
            "by_status": by_status,
            "processing": len([t for t in self._processing_tasks.values() if not t.done()]),
        }


# Global uploader
_global_uploader: Optional[KnowledgeUploader] = None


def get_knowledge_uploader() -> KnowledgeUploader:
    """Get or create global knowledge uploader."""
    global _global_uploader
    if _global_uploader is None:
        _global_uploader = KnowledgeUploader()
    return _global_uploader
